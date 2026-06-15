"""Turn the designed Letter_of_Offer.docx into a docxtemplater template.
Scalars become {tokens}; the Guarantor and Security tables keep their fixed 8 rows
(positional tokens, original formatting preserved); every table is locked so it can't
split across a page break."""
import docx, re, sys
from copy import deepcopy
from docx.text.paragraph import Paragraph
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

SRC = "templates/source_Letter_of_Offer.docx"
OUT = "templates/letter-of-offer.template.docx"

# --- scalar replacements: literal text in a run -> token (formatting preserved) ---
SCALARS = {
    "[BORROWER ENTITY NAME] ACN [XXX XXX XXX]": "{borrower.entity_name} ACN {borrower.acn}",
    "[PRODUCT NAME]": "{deal.product}",
    "[INTRODUCER NAME]": "{broker.name}",
    "[COMPANY NAME] ABN [XX XXX XXX XXX]": "{broker.company_name} ABN {broker.company_abn}",
    "Greenlink Capital Pty Ltd is pleased to set out the indicative terms on which it is prepared to provide a loan facility to [BORROWER ENTITY NAME], as detailed in this Letter of Offer.":
        "Greenlink Capital Pty Ltd is pleased to set out the indicative terms on which it is prepared to provide a loan facility to {borrower.entity_name}, as detailed in this Letter of Offer.",
    "[XX] months from settlement": None,  # handled positionally below
    "Early repayment permitted [at any time / after X months] subject to a minimum of [X] months\u2019 interest":
        "Early repayment permitted after the minimum interest term, subject to a minimum of {facility.minimum_interest_months} months\u2019 interest",
    "[INTEREST ONLY / PRINCIPAL AMD INTEREST/ INTEREST CAPITALISED]": "{facility.repayment_type}",
    "[LOAN PURPOSE]": "{facility.loan_purpose}",
    "[EXIT STRATEGY]": "{facility.exit_strategy}",
    "$XXXXX": "${facility.loan_amount}",
    "XX.XX% p.a.": "{facility.interest_rate}% p.a.",
    "[BROKER]": "{broker.name}",
    "[PANEL SOLICITOR]": "{deal.panel_solicitor}",
    "[GLC-INV-XXXX]": "{invoice.number}",
    "[GLC-XXXX]": "{deal.reference}",
    " xxx-xxx-xxx      ": " {deal.reference}      ",
    " xxxx days": " {deal.validity_days} days",
    "[BDM NAME], Business Development \u2014 ": "{bdm.name}, Business Development \u2014 ",
    "[MOBILE] \u00b7 [BDM EMAIL]": "{bdm.mobile} \u00b7 {bdm.email}",
    " xxxx": " {deal.loo_date}",
}

# placeholders split across bold runs -> applied by flattening SHORT cell paragraphs only
SPLIT = [
    ("[COMPANY NAME] ABN [XX XXX XXX XXX]", "{broker.company_name} ABN {broker.company_abn}"),
    ("[ENTITY NAME] ACN [XXX XXX XXX]", "{borrower.entity_name} ACN {borrower.acn}"),
    ("[BORROWER ENTITY NAME] ACN [XXX XXX XXX]", "{borrower.entity_name} ACN {borrower.acn}"),
    ("[PRODUCT NAME]", "{deal.product}"),
    ("[COMPANY NAME]", "{broker.company_name}"),
    ("[BORROWER ENTITY NAME]", "{borrower.entity_name}"),
    ("[ENTITY NAME]", "{borrower.entity_name}"),
    ("[XX XXX XXX XXX]", "{borrower.abn}"),
    ("[XXX XXX XXX]", "{borrower.acn}"),
]
def merge_replace(paragraph):
    full = "".join(r.text for r in paragraph.runs)
    if "[" not in full or len(full) > 90:
        return
    new = full
    for k, v in SPLIT:
        new = new.replace(k, v)
    if new != full and paragraph.runs:
        paragraph.runs[0].text = new
        for r in paragraph.runs[1:]:
            r.text = ""
def global_run_fix(paragraph):
    for r in paragraph.runs:
        if "[ENTITY NAME]" in r.text:
            r.text = r.text.replace("[ENTITY NAME]", "{borrower.entity_name}")

def repl_runs_in_paragraph(p):
    # join runs, replace, write back into first run if a known scalar is fully inside one run
    for r in p.runs:
        for k, v in SCALARS.items():
            if v is not None and k in r.text:
                r.text = r.text.replace(k, v)

def set_cell(cell, text):
    # replace cell content with one run, keeping the cell's first run formatting if any
    if cell.paragraphs and cell.paragraphs[0].runs:
        base = cell.paragraphs[0].runs[0]
        base.text = text
        for extra in cell.paragraphs[0].runs[1:]:
            extra.text = ""
        for extrap in cell.paragraphs[0:0]:
            pass
    else:
        cell.text = text

def cell_text(cell):
    return cell.text.strip()

def first(cell):
    return cell.text.strip().upper()

def fill_fixed_rows(tbl, colmap, prefix):
    """Put positional tokens ({prefix}{rownum}_{field}) into the data cells of every
    numbered row, cloning row 1's paragraph formatting so blank rows match. Deletes nothing."""
    fmt = {ci: deepcopy(tbl.rows[1].cells[ci].paragraphs[0]._p) for ci in colmap}
    rn = 0
    for row in tbl.rows[1:]:
        if "TOTAL" in row.cells[0].text.upper():
            continue
        rn += 1
        for ci, field in colmap.items():
            token = "{" + f"{prefix}{rn}_{field}" + "}"
            dst = row.cells[ci]
            tc = dst._tc
            for p in list(dst.paragraphs):
                tc.remove(p._p)
            new_p = deepcopy(fmt[ci])
            tc.append(new_p)
            para = Paragraph(new_p, dst)
            if para.runs:
                para.runs[0].text = token
                for r in para.runs[1:]:
                    r.text = ""
            else:
                para.add_run(token)

def lock_table(tbl):
    """cantSplit on every row + keepNext within the table so it never breaks across a page."""
    rows = tbl.rows
    for i, row in enumerate(rows):
        trPr = row._tr.get_or_add_trPr()
        if trPr.find(qn("w:cantSplit")) is None:
            trPr.append(OxmlElement("w:cantSplit"))
        if i < len(rows) - 1:  # not the last row (avoid pulling following content up)
            for cell in row.cells:
                for p in cell.paragraphs:
                    pPr = p._p.get_or_add_pPr()
                    if pPr.find(qn("w:keepNext")) is None:
                        pPr.append(OxmlElement("w:keepNext"))

# exact single-run replacements (handles split placeholders & duplicates safely)
RUN_EXACT = {
    "Dear [": "Dear ",
    "BORROWER ENTITY NAME]": "{borrower.entity_name}",
    "[ENTITY NAME] ACN [XXX XXX XXX]": "{borrower.entity_name} ACN {borrower.acn}",
    "[DATE]": "{invoice.issue_date}",
    "[XX XXX XXX XXX]": "{borrower.abn}",
    "$85,250.00": "${funding.total_costs}",
    "$2,364,750.00": "${funding.funds_available}",
    "$2,693,775.00": "${funding.total_repayment}",
}

doc = docx.Document(SRC)

# 1) paragraph-level scalars
for p in doc.paragraphs:
    repl_runs_in_paragraph(p)
    for r in p.runs:
        if r.text in RUN_EXACT:
            r.text = RUN_EXACT[r.text]
    merge_replace(p)
    global_run_fix(p)

# helper to find a table by a header signature
def header_sig(tbl):
    try:
        return " | ".join(c.text.strip().upper() for c in tbl.rows[0].cells)
    except Exception:
        return ""

GUAR_DONE = SEC_DONE = False
DATE_SLOTS = []  # collect the two "[XX] months from settlement" + facility-terms cells

for tbl in doc.tables:
    sig = header_sig(tbl)

    # run scalar + exact replace inside every cell paragraph first
    for row in tbl.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                repl_runs_in_paragraph(p)
                for r in p.runs:
                    if r.text in RUN_EXACT:
                        r.text = RUN_EXACT[r.text]
                merge_replace(p)
                global_run_fix(p)

    # ---- TRANSACTION SUMMARY: LOAN AMOUNT | INTEREST RATE | LVR | TERM | LOAN TYPE ----
    if "LOAN AMOUNT" in sig and "LVR" in sig and "TERM" in sig:
        d = tbl.rows[1].cells
        set_cell(d[0], "${facility.loan_amount}")
        set_cell(d[1], "{facility.interest_rate}% p.a.")
        set_cell(d[2], "{facility.lvr}%")
        set_cell(d[3], "{facility.term_months} months")
        set_cell(d[4], "{facility.facility_type}")
        continue

    # ---- FACILITY TERMS / DRAWDOWN / INVOICE: label|value rows ----
    label_map = {
        "LOAN TERM": ("{facility.term_months} months from settlement", 1),
        "PRE-PAID INTEREST TERM": ("{facility.prepaid_interest_months} months from settlement", 1),
        "LOAN FACILITY AMOUNT": ("${drawdown.facility_amount}", 1),
        "NET SURPLUS TO BORROWER AT SETTLEMENT": ("${drawdown.net_surplus}", 1),
    }
    handled = False
    for row in tbl.rows:
        lab = first(row.cells[0])
        # drawdown 'less' lines (label varies, match on keyword)
        if lab.startswith("LESS") and "PREPAID INTEREST" in lab:
            set_cell(row.cells[1], "${drawdown.less_prepaid_interest}"); handled = True
        elif lab.startswith("LESS") and "MONTHLY BORROWING" in lab:
            set_cell(row.cells[1], "${drawdown.less_prepaid_mbr}"); handled = True
        elif lab.startswith("LESS") and "ESTABLISHMENT" in lab:
            set_cell(row.cells[1], "${drawdown.less_establishment}"); handled = True
        elif lab.startswith("LESS") and "BROKERAGE" in lab:
            set_cell(row.cells[1], "${drawdown.less_brokerage}"); handled = True
        elif lab.startswith("LESS") and ("LEGAL" in lab or "DOCUMENTATION" in lab):
            set_cell(row.cells[1], "${drawdown.less_legal}"); handled = True
        elif lab in label_map:
            txt, ci = label_map[lab]; set_cell(row.cells[ci], txt); handled = True

    # ---- FUNDING SCHEDULE: ITEM | RATE | AMOUNT ----
    if "ITEM" in sig and "RATE" in sig and "AMOUNT" in sig:
        for row in tbl.rows[1:]:
            lab = first(row.cells[0])
            if "FIRST MORTGAGE INTEREST" in lab:
                set_cell(row.cells[1], "{funding.interest_rate}% p.a."); set_cell(row.cells[2], "${funding.interest_amount}")
            elif "MONTHLY BORROWING" in lab:
                set_cell(row.cells[1], "{funding.mbr_rate}% p.m."); set_cell(row.cells[2], "${funding.mbr_amount}")
            elif "ESTABLISHMENT" in lab:
                set_cell(row.cells[1], "{funding.estab_rate}%"); set_cell(row.cells[2], "${funding.estab_amount}")
            elif lab == "BROKERAGE":
                set_cell(row.cells[1], "{funding.brokerage_rate}%"); set_cell(row.cells[2], "${funding.brokerage_amount}")
            elif "LEGAL" in lab:
                set_cell(row.cells[2], "${funding.legal_fees}")
        continue

    # ---- INVOICE LINE ITEMS: ITEM | AMOUNT | GST | TOTAL ----
    if "ITEM" in sig and "GST" in sig and "TOTAL" in sig:
        for row in tbl.rows[1:]:
            lab = first(row.cells[0])
            if "DUE DILIGENCE" in lab:
                set_cell(row.cells[1], "$ {invoice.dd_fee}"); set_cell(row.cells[2], "$ {invoice.dd_gst}"); set_cell(row.cells[3], "$ {invoice.dd_total}")
            elif "VALUATION" in lab:
                set_cell(row.cells[1], "$ {invoice.val_fee}"); set_cell(row.cells[2], "$ {invoice.val_gst}"); set_cell(row.cells[3], "$ {invoice.val_total}")
            elif "LEGAL FEE" in lab:
                set_cell(row.cells[1], "$ {invoice.legal_fee}"); set_cell(row.cells[2], "$ {invoice.legal_gst}"); set_cell(row.cells[3], "$ {invoice.legal_total}")
            elif "TOTAL DUE" in lab:
                set_cell(row.cells[-1], "$ {invoice.total_due}")
            elif "TOTAL EXCLUDING GST" in lab:
                set_cell(row.cells[-1], "$ {invoice.total_excl_gst}")
            elif "TOTAL INCLUDING GST" in lab:
                set_cell(row.cells[-1], "$ {invoice.total_incl_gst}")
            elif "ALREADY PAID" in lab:
                set_cell(row.cells[-1], "$ {invoice.already_paid}")
        continue

    if handled:
        continue

    # ---- GUARANTOR SCHEDULE: # | GUARANTOR | TYPE | CAPACITY | ILA (8 fixed rows) ----
    if (not GUAR_DONE) and "GUARANTOR" in sig and "CAPACITY" in sig and "ILA" in sig:
        fill_fixed_rows(tbl, {1: "name", 2: "type", 3: "capacity", 4: "ila"}, "g")
        GUAR_DONE = True
        continue

    # ---- SECURITY SCHEDULE: # | SECURITY ... | MORTGAGE | LVR | ESTIMATED VALUE + TOTAL row ----
    if (not SEC_DONE) and "MORTGAGE" in sig and "ESTIMATED VALUE" in sig:
        fill_fixed_rows(tbl, {1: "address", 2: "mortgage", 3: "lvr", 4: "value"}, "s")
        for row in tbl.rows:
            if "TOTAL" in row.cells[0].text.upper():
                set_cell(row.cells[-1], "{facility.total_security_value}")
                break
        SEC_DONE = True
        continue

for tbl in doc.tables:
    lock_table(tbl)

doc.save(OUT)
print("GUAR_DONE", GUAR_DONE, "SEC_DONE", SEC_DONE)
print("saved", OUT)
